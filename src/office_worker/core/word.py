"""Generación de Word (.docx) con python-docx, aplicando tema corporativo."""
from __future__ import annotations
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from .security import safe_out

def _hex_to_rgb(h: str) -> RGBColor:
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

def _clean_font_name(font_str: str | None, default: str = "Arial") -> str:
    if not font_str:
        return default
    first = str(font_str).split(",")[0].strip().strip("'").strip('"')
    return first or default

def _format_agency_table(tbl, headers: list, rows: list, th: dict) -> None:
    """Aplica formato de tabla corporativo-agencia con cabecera rellena, zebra y bordes sutiles."""
    primary_hex = th["primary"].lstrip("#")
    row_alt_hex = th["row_alt"].lstrip("#")
    font_title = _clean_font_name(th.get("font_title"), "Helvetica")
    font_body = _clean_font_name(th.get("font_body"), "Arial")
    text_color = _hex_to_rgb(th["text"])
    white_color = RGBColor(255, 255, 255)

    tbl.style = "Normal Table"
    tbl.autofit = False

    # Margen interno cómodo para celdas (padding) y bordes sutiles
    tblPr = tbl._tbl.tblPr
    tblPr.append(parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        f'<w:top w:w="140" w:type="dxa"/>'
        f'<w:bottom w:w="140" w:type="dxa"/>'
        f'<w:left w:w="180" w:type="dxa"/>'
        f'<w:right w:w="180" w:type="dxa"/>'
        f'</w:tblCellMar>'
    ))
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>'
        f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    ))

    n_cols = len(headers) if headers else (len(rows[0]) if rows else 1)
    col_w = Inches(6.9 / max(n_cols, 1))

    # Cabecera de la tabla
    if headers:
        hdr_row = tbl.rows[0]
        hdr_row._tr.get_or_add_trPr().append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
        hdr_row._tr.get_or_add_trPr().append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        for i, hv in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.width = col_w
            cell.text = str(hv).upper()
            cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{primary_hex}"/>'))
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            for r in p.runs:
                r.font.name = font_title
                r.font.size = Pt(8.5)
                r.font.bold = True
                r.font.color.rgb = white_color
                r._r.get_or_add_rPr().append(parse_xml(f'<w:spacing {nsdecls("w")} w:val="18"/>'))

    # Filas de datos
    start_r = 1 if headers else 0
    for r_idx, row in enumerate(rows):
        tbl_row = tbl.rows[r_idx + start_r]
        tbl_row._tr.get_or_add_trPr().append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        is_alt = (r_idx % 2 == 1)
        for c_idx, cv in enumerate(row):
            if c_idx < len(tbl_row.cells):
                cell = tbl_row.cells[c_idx]
                cell.width = col_w
                cell.text = str(cv)
                tcPr = cell._tc.get_or_add_tcPr()
                if is_alt:
                    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{row_alt_hex}"/>'))
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.15
                for r in p.runs:
                    r.font.name = font_body
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = text_color
                    if c_idx == 0:
                        r.font.bold = True

def _add_callout(doc, text: str, th: dict, title: str | None = None) -> None:
    """Inserta un bloque de callout/card con fondo row_alt y barra izquierda accent."""
    accent_hex = th["accent"].lstrip("#")
    row_alt_hex = th["row_alt"].lstrip("#")
    text_color = _hex_to_rgb(th["text"])
    font_body = _clean_font_name(th.get("font_body"), "Arial")
    font_title = _clean_font_name(th.get("font_title"), "Helvetica")

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Normal Table"
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.9)

    tblPr = tbl._tbl.tblPr
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{accent_hex}"/>'
        f'<w:top w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:insideH w:val="none"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    ))
    tblPr.append(parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        f'<w:top w:w="160" w:type="dxa"/>'
        f'<w:bottom w:w="160" w:type="dxa"/>'
        f'<w:left w:w="200" w:type="dxa"/>'
        f'<w:right w:w="200" w:type="dxa"/>'
        f'</w:tblCellMar>'
    ))

    cell = tbl.cell(0, 0)
    cell.width = Inches(6.9)
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{row_alt_hex}"/>'))

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2

    if title:
        rt = p.add_run(f"{title}\n")
        rt.font.name = font_title
        rt.font.bold = True
        rt.font.size = Pt(10)
        rt.font.color.rgb = _hex_to_rgb(th["primary"])

    r = p.add_run(str(text))
    r.font.name = font_body
    r.font.size = Pt(9.5)
    r.font.color.rgb = text_color

    # Espacio tras el callout
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(6)

def create_word(out_path, title="", subtitle=None, blocks=None, theme=None, template_docx=None, context=None, **kwargs):
    """Crea un .docx profesional desde bloques declarativos o rellenando una plantilla con docxtpl.

    - blocks: lista de dicts. Tipos soportados:
      {"type":"h1"|"h2"|"h3"|"subheading"|"p", "text": "..."}
      {"type":"kicker", "text": "..."}
      {"type":"callout"|"card", "text": "...", "title"?: "..."}
      {"type":"bullet_list", "items": ["...", ...]}
      {"type":"table", "headers":[...], "rows":[[...], ...]}
    - template_docx: ruta opcional a plantilla .docx existente con placeholders Jinja {{ variable }}.
    - context: diccionario opcional de variables para rellenar template_docx.
    Si blocks es None y no hay template_docx → documento mínimo con portada/título. Devuelve ruta absoluta.
    """
    out_path = safe_out(out_path)

    if template_docx:
        from .templates_pack import resolve_template_path
        tpl_path = resolve_template_path(template_docx)
        from docxtpl import DocxTemplate
        tpl = DocxTemplate(tpl_path)
        ctx = dict(context or {})
        if title and "title" not in ctx: ctx["title"] = title
        if subtitle and "subtitle" not in ctx: ctx["subtitle"] = subtitle
        tpl.render(ctx)
        tpl.save(out_path)
        if os.path.getsize(out_path) < 300:
            raise RuntimeError(f"DOCX generado sospechosamente pequeño ({os.path.getsize(out_path)}B)")
        return out_path

    from .themes import load_theme
    th = load_theme(theme)
    primary = _hex_to_rgb(th["primary"])
    accent = _hex_to_rgb(th["accent"])
    text_color = _hex_to_rgb(th["text"])
    muted_color = _hex_to_rgb(th["muted"])
    primary_hex = th["primary"].lstrip("#")
    accent_hex = th["accent"].lstrip("#")

    font_title = _clean_font_name(th.get("font_title"), "Helvetica")
    font_body = _clean_font_name(th.get("font_body"), "Arial")

    doc = Document()

    # Geometría de página editorial (márgenes consistentes ~2cm)
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.95)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Estilo base Normal según tema
    st = doc.styles["Normal"]
    st.font.name = font_body
    st.font.size = Pt(10.5)
    st.font.color.rgb = text_color

    # Determinar si el primer bloque es un kicker que acompaña a la portada
    blocks_list = list(blocks or [])
    cover_kicker = kwargs.get("kicker")
    if not cover_kicker and blocks_list and blocks_list[0].get("type") == "kicker" and title:
        cover_kicker = blocks_list[0].get("text")
        blocks_list = blocks_list[1:]

    date_text = kwargs.get("date")

    # --- PORTADA EJECUTIVA (cuando hay title) ---
    if title:
        # Banda superior de color primary
        p_band = doc.add_paragraph()
        p_band.paragraph_format.space_before = Pt(0)
        p_band.paragraph_format.space_after = Pt(12)
        p_bandPr = p_band._p.get_or_add_pPr()
        p_bandPr.append(parse_xml(
            f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="36" w:space="1" w:color="{primary_hex}"/><w:top w:val="none"/><w:left w:val="none"/><w:right w:val="none"/></w:pBdr>'
        ))

        # Kicker en accent uppercase letter-spacing
        if cover_kicker:
            p_k = doc.add_paragraph()
            p_k.paragraph_format.space_before = Pt(0)
            p_k.paragraph_format.space_after = Pt(4)
            p_k.paragraph_format.keep_with_next = True
            rk = p_k.add_run(str(cover_kicker).upper())
            rk.font.name = font_title
            rk.font.size = Pt(9)
            rk.font.bold = True
            rk.font.color.rgb = accent
            rk._r.get_or_add_rPr().append(parse_xml(f'<w:spacing {nsdecls("w")} w:val="24"/>'))

        # Título grande (26pt bold, color primary)
        h = doc.add_heading(title, level=0)
        h.paragraph_format.space_before = Pt(2)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        # Quitar el borde inferior por defecto del estilo Title de Word
        h._p.get_or_add_pPr().append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="none"/></w:pBdr>'))
        for r in h.runs:
            r.font.name = font_title
            r.font.size = Pt(26)
            r.font.bold = True
            r.font.color.rgb = primary

        # Subtítulo en muted
        if subtitle:
            p_sub = doc.add_paragraph(subtitle)
            p_sub.paragraph_format.space_before = Pt(0)
            p_sub.paragraph_format.space_after = Pt(4)
            p_sub.paragraph_format.keep_with_next = True
            for r in p_sub.runs:
                r.font.name = font_body
                r.font.size = Pt(11)
                r.font.color.rgb = muted_color

        # Fecha opcional en muted
        if date_text:
            p_dt = doc.add_paragraph()
            p_dt.paragraph_format.space_before = Pt(0)
            p_dt.paragraph_format.space_after = Pt(6)
            p_dt.paragraph_format.keep_with_next = True
            rd = p_dt.add_run(str(date_text))
            rd.font.name = font_body
            rd.font.size = Pt(9.5)
            rd.font.color.rgb = muted_color

        # Separador fino bajo el bloque de portada
        p_sep = doc.add_paragraph()
        p_sep.paragraph_format.space_before = Pt(4)
        p_sep.paragraph_format.space_after = Pt(18)
        p_sepPr = p_sep._p.get_or_add_pPr()
        p_sepPr.append(parse_xml(
            f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CBD5E1"/><w:top w:val="none"/><w:left w:val="none"/><w:right w:val="none"/></w:pBdr>'
        ))

    # --- CUERPO DEL DOCUMENTO ---
    for b in blocks_list:
        t = b.get("type", "p")

        if t == "h1":
            hh = doc.add_heading(b.get("text", ""), level=1)
            hh.paragraph_format.space_before = Pt(20)
            hh.paragraph_format.space_after = Pt(8)
            hh.paragraph_format.keep_with_next = True
            for r in hh.runs:
                r.font.name = font_title
                r.font.size = Pt(18)
                r.font.bold = True
                r.font.color.rgb = primary
            hh._p.get_or_add_pPr().append(parse_xml(
                f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="18" w:space="6" w:color="{primary_hex}"/><w:left w:val="none"/><w:right w:val="none"/><w:top w:val="none"/></w:pBdr>'
            ))

        elif t == "h2":
            hh = doc.add_heading(b.get("text", ""), level=2)
            hh.paragraph_format.space_before = Pt(16)
            hh.paragraph_format.space_after = Pt(6)
            hh.paragraph_format.keep_with_next = True
            for r in hh.runs:
                r.font.name = font_title
                r.font.size = Pt(13.5)
                r.font.bold = True
                r.font.color.rgb = primary
            hh._p.get_or_add_pPr().append(parse_xml(
                f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="8" w:color="{accent_hex}"/><w:bottom w:val="none"/><w:top w:val="none"/><w:right w:val="none"/></w:pBdr>'
            ))

        elif t in ("h3", "subheading"):
            hh = doc.add_heading(b.get("text", ""), level=3)
            hh.paragraph_format.space_before = Pt(12)
            hh.paragraph_format.space_after = Pt(4)
            hh.paragraph_format.keep_with_next = True
            for r in hh.runs:
                r.font.name = font_title
                r.font.size = Pt(11)
                r.font.bold = True
                r.font.color.rgb = primary

        elif t == "kicker":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            rk = p.add_run(str(b.get("text", "")).upper())
            rk.font.name = font_title
            rk.font.size = Pt(8.5)
            rk.font.bold = True
            rk.font.color.rgb = accent
            rk._r.get_or_add_rPr().append(parse_xml(f'<w:spacing {nsdecls("w")} w:val="24"/>'))

        elif t in ("callout", "card"):
            _add_callout(doc, text=b.get("text") or b.get("content") or "", th=th, title=b.get("title"))

        elif t == "bullet_list":
            items = b.get("items") or b.get("bullets") or []
            if isinstance(items, str):
                items = [items]
            for item in items:
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_before = Pt(2)
                bp.paragraph_format.space_after = Pt(2)
                bp.paragraph_format.line_spacing = 1.15
                r = bp.add_run(str(item))
                r.font.name = font_body
                r.font.size = Pt(10.5)
                r.font.color.rgb = text_color

        elif t == "table":
            headers = b.get("headers", [])
            rows = b.get("rows", [])
            n_rows = len(rows) + (1 if headers else 0)
            n_cols = len(headers) if headers else (len(rows[0]) if rows else 1)
            tbl = doc.add_table(rows=n_rows, cols=n_cols)
            _format_agency_table(tbl, headers, rows, th)
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(0)
            sp.paragraph_format.space_after = Pt(8)

        else:  # p / default
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.2
            p.paragraph_format.widow_control = True
            r = p.add_run(b.get("text", ""))
            r.font.name = font_body
            r.font.size = Pt(10.5)
            r.font.color.rgb = text_color
            if b.get("bold"):
                r.bold = True
            if b.get("italic"):
                r.italic = True

    doc.save(out_path)
    if os.path.getsize(out_path) < 300:
        raise RuntimeError(f"DOCX sospechosamente pequeño ({os.path.getsize(out_path)}B)")
    return out_path


def edit_word(
    input_path: str,
    operations: list[dict[str, Any]] | str,
    output_path: str | None = None,
) -> dict[str, Any]:
    """Modifica un documento Word existente (.docx) preservando estilos y formato.

    Operaciones soportadas (lista de dicts):
    - {"op": "append_paragraph", "text": str, "style"?: str, "bold"?: bool, "italic"?: bool}
    - {"op": "replace_text", "find"|"old": str, "replace"|"new": str, "count"?: int}
    - {"op": "insert_after_heading", "heading_text": str, "text": str, "style"?: str, "bold"?: bool, "italic"?: bool}
    - {"op": "append_table", "headers": list, "rows": list[list], "style"?: str}

    Devuelve dict con fidelity honesta ("clean"), warnings y ruta absoluta.
    """
    import json

    input_path = os.path.abspath(os.path.expanduser(str(input_path)))
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Archivo Word no encontrado: {input_path}")

    target_out = safe_out(output_path if output_path else input_path)
    doc = Document(input_path)

    warnings = ["python-docx edits document body; embedded OLE objects, drawing canvas and SmartArt are not updated."]
    fidelity = "clean"

    if isinstance(operations, str):
        try:
            ops = json.loads(operations)
        except json.JSONDecodeError as exc:
            raise ValueError(f"operations inválido (JSON malformado): {exc}")
    else:
        ops = list(operations or [])

    for op in ops:
        op_name = str(op.get("op") or op.get("operation") or "").lower().strip()

        if op_name == "append_paragraph":
            text = op.get("text", "")
            style = op.get("style")
            p = doc.add_paragraph(text, style=style)
            bold = op.get("bold")
            italic = op.get("italic")
            if bold is not None or italic is not None:
                for r in p.runs:
                    if bold is not None:
                        r.bold = bool(bold)
                    if italic is not None:
                        r.italic = bool(italic)

        elif op_name == "replace_text":
            old = op.get("find") or op.get("old") or ""
            new = op.get("replace") or op.get("new") or ""
            count = op.get("count")
            max_c = int(count) if count is not None else -1
            c_left = max_c

            if not old:
                continue

            # Reemplazar en párrafos
            for p in doc.paragraphs:
                if old in p.text:
                    replaced_in_run = False
                    for r in p.runs:
                        if old in r.text:
                            if c_left > 0:
                                r.text = r.text.replace(old, new, c_left)
                                c_left -= 1
                            elif c_left == -1:
                                r.text = r.text.replace(old, new)
                            replaced_in_run = True
                            if c_left == 0:
                                break
                    if not replaced_in_run:
                        if c_left > 0:
                            p.text = p.text.replace(old, new, c_left)
                            c_left -= 1
                        elif c_left == -1:
                            p.text = p.text.replace(old, new)
                    if c_left == 0:
                        break

            # Reemplazar en tablas
            if c_left != 0:
                for tbl in doc.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if old in p.text:
                                    if c_left > 0:
                                        p.text = p.text.replace(old, new, c_left)
                                        c_left -= 1
                                    elif c_left == -1:
                                        p.text = p.text.replace(old, new)
                                    if c_left == 0:
                                        break

        elif op_name == "insert_after_heading":
            target_heading = str(op.get("heading_text", "")).strip().lower()
            text = op.get("text", "")
            style = op.get("style", "Normal")
            matched_p = None
            for p in doc.paragraphs:
                if target_heading in p.text.strip().lower():
                    matched_p = p
                    break

            if matched_p is not None:
                new_p = doc.add_paragraph(text, style=style)
                bold = op.get("bold")
                italic = op.get("italic")
                if bold is not None or italic is not None:
                    for r in new_p.runs:
                        if bold is not None:
                            r.bold = bool(bold)
                        if italic is not None:
                            r.italic = bool(italic)
                # Mover XML para insertar inmediatamente después del encabezado
                matched_p._p.addnext(new_p._p)
            else:
                # Si no encontró el heading, hacer append como fallback con warning
                p = doc.add_paragraph(text, style=style)
                warnings.append(f"Heading '{op.get('heading_text')}' not found; appended paragraph at end.")

        elif op_name == "append_table":
            headers = op.get("headers", [])
            rows = op.get("rows", [])
            tbl_style = op.get("style", "Light Grid Accent 1")
            tbl = doc.add_table(rows=1, cols=len(headers))
            tbl.style = tbl_style
            hdr = tbl.rows[0].cells
            for i, hv in enumerate(headers):
                hdr[i].text = str(hv)
            for row in rows:
                cells = tbl.add_row().cells
                for i, cv in enumerate(row):
                    if i < len(cells):
                        cells[i].text = str(cv)

        else:
            raise ValueError(f"Operación de edición no soportada: '{op_name}'")

    doc.save(target_out)
    return {
        "status": "ok",
        "path": os.path.abspath(target_out),
        "bytes": os.path.getsize(target_out),
        "fidelity": fidelity,
        "warnings": warnings,
        "operations_applied": len(ops),
    }


def mail_merge(
    template_path: str,
    dataset_csv: str = "",
    dataset_json: str = "",
    output_prefix: str = "",
    fields: list[str] | str | None = None,
) -> dict[str, Any]:
    """Genera N documentos .docx combinando una plantilla docxtpl con un dataset CSV o JSON.

    - template_path: ruta a la plantilla .docx con placeholders {{ variable }}.
    - dataset_csv: ruta a archivo CSV o texto CSV plano.
    - dataset_json: ruta a archivo JSON o texto/estructura JSON.
    - output_prefix: prefijo para los archivos generados (ej: 'docs/carta' -> 'docs/carta_1.docx').
    - fields: lista o nombres de campos opcionales para filtrar columnas.

    Devuelve dict con status, n_docs, paths y fields.
    """
    import io
    import json
    from pathlib import Path
    import pandas as pd
    from docxtpl import DocxTemplate
    from .templates_pack import resolve_template_path

    try:
        tpl_path = resolve_template_path(template_path)
    except Exception:
        tpl_path = os.path.abspath(os.path.expanduser(str(template_path)))

    if not os.path.exists(tpl_path):
        raise FileNotFoundError(f"Plantilla no encontrada: {template_path}")

    df = None
    if dataset_csv:
        csv_str = str(dataset_csv).strip()
        if os.path.isfile(csv_str):
            df = pd.read_csv(csv_str)
        else:
            df = pd.read_csv(io.StringIO(csv_str))
    elif dataset_json:
        if isinstance(dataset_json, str):
            json_str = dataset_json.strip()
            if os.path.isfile(json_str):
                with open(json_str, "r", encoding="utf-8") as f:
                    raw_data = json.load(f)
            else:
                raw_data = json.loads(json_str)
        else:
            raw_data = dataset_json

        if isinstance(raw_data, list):
            df = pd.DataFrame(raw_data)
        elif isinstance(raw_data, dict):
            rows = raw_data.get("rows") or raw_data.get("data") or [raw_data]
            df = pd.DataFrame(rows)
        else:
            raise ValueError(f"Formato JSON no reconocido para dataset: {type(raw_data)}")
    else:
        raise ValueError("Debe especificarse dataset_csv o dataset_json.")

    if df is None or df.empty:
        raise ValueError("El dataset está vacío o no contiene filas.")

    if fields:
        if isinstance(fields, str):
            try:
                f_list = json.loads(fields)
            except Exception:
                f_list = [f.strip() for f in fields.split(",") if f.strip()]
        else:
            f_list = list(fields)
        valid_cols = [c for c in f_list if c in df.columns]
        if valid_cols:
            df = df[valid_cols]

    df.columns = [str(c) for c in df.columns]
    records = df.fillna("").to_dict(orient="records")

    if not output_prefix:
        base_dir = os.path.dirname(os.path.abspath(tpl_path))
        stem = Path(tpl_path).stem
        output_prefix = os.path.join(base_dir, f"{stem}_merged")
    else:
        output_prefix = str(output_prefix)
        if output_prefix.endswith(".docx"):
            output_prefix = output_prefix[:-5]

    paths = []
    for idx, rec in enumerate(records, start=1):
        target_file = safe_out(f"{output_prefix}_{idx}.docx")
        tpl = DocxTemplate(tpl_path)
        tpl.render(rec)
        tpl.save(target_file)
        if os.path.getsize(target_file) < 300:
            raise RuntimeError(f"DOCX generado sospechosamente pequeño ({os.path.getsize(target_file)}B): {target_file}")
        paths.append(os.path.abspath(target_file))

    return {
        "status": "ok",
        "template": os.path.abspath(tpl_path),
        "n_docs": len(paths),
        "paths": paths,
        "fields": list(df.columns),
    }

