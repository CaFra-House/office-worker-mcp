"""Comparación y diff honesto entre documentos Word (.docx) y PDF."""
from __future__ import annotations
import difflib
import os
from pathlib import Path
from typing import Any

from .office_reader import read_office


def _extract_document_paragraphs(path: str) -> list[str]:
    """Extrae párrafos y elementos textuales de un documento (.docx, .pdf, o texto)."""
    ext = Path(path).suffix.lower()

    if ext == ".docx":
        from docx import Document
        doc = Document(path)
        paragraphs = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                paragraphs.append(t)
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append("[Tabla] " + " | ".join(cells))
        return paragraphs

    if ext == ".pdf":
        import fitz
        doc = fitz.open(path)
        paragraphs = []
        for page in doc:
            page_text = page.get_text("text")
            for block in page_text.split("\n\n"):
                cleaned = " ".join(block.split())
                if cleaned:
                    paragraphs.append(cleaned)
        doc.close()
        return paragraphs

    if ext in (".pptx", ".xlsx", ".xlsm"):
        data = read_office(path, format="json")
        paragraphs = []
        if ext == ".pptx":
            for s in data.get("slides", []):
                for t in s.get("text", []):
                    if t.strip():
                        paragraphs.append(t.strip())
        else:
            for sh in data.get("sheets", []):
                for r in sh.get("rows", []):
                    line = " | ".join(str(c) for c in r if str(c).strip())
                    if line:
                        paragraphs.append(line)
        return paragraphs

    # Archivo de texto plano fallback
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return [line.strip() for line in f if line.strip()]


def document_diff(
    path_a: str,
    path_b: str,
    format: str = "json",
) -> dict[str, Any]:
    """Compara dos documentos Word (.docx) o PDF y retorna diferencias textuales honestas.

    - path_a: ruta al documento base (versión anterior/original).
    - path_b: ruta al documento modificado (versión nueva/revisada).
    - format: 'json' (default) o 'markdown'.

    Usa difflib sobre el texto extraído.
    Retorna advertencia explícita indicando que es un diff textual y no un redline legal-grade.
    """
    path_a = os.path.abspath(os.path.expanduser(str(path_a)))
    path_b = os.path.abspath(os.path.expanduser(str(path_b)))

    if not os.path.exists(path_a):
        raise FileNotFoundError(f"Documento A no encontrado: {path_a}")
    if not os.path.exists(path_b):
        raise FileNotFoundError(f"Documento B no encontrado: {path_b}")

    paras_a = _extract_document_paragraphs(path_a)
    paras_b = _extract_document_paragraphs(path_b)

    matcher = difflib.SequenceMatcher(None, paras_a, paras_b)
    opcodes = matcher.get_opcodes()

    diffs = []
    count_added = 0
    count_deleted = 0
    count_modified = 0
    count_unchanged = 0

    for tag, i1, i2, j1, j2 in opcodes:
        if tag == "equal":
            count_unchanged += (i2 - i1)
        elif tag == "insert":
            for j in range(j1, j2):
                diffs.append({"type": "added", "text": paras_b[j]})
                count_added += 1
        elif tag == "delete":
            for i in range(i1, i2):
                diffs.append({"type": "deleted", "text": paras_a[i]})
                count_deleted += 1
        elif tag == "replace":
            sub_a = paras_a[i1:i2]
            sub_b = paras_b[j1:j2]
            min_len = min(len(sub_a), len(sub_b))
            for k in range(min_len):
                diffs.append({
                    "type": "modified",
                    "old_text": sub_a[k],
                    "new_text": sub_b[k],
                })
                count_modified += 1
            if len(sub_a) > min_len:
                for k in range(min_len, len(sub_a)):
                    diffs.append({"type": "deleted", "text": sub_a[k]})
                    count_deleted += 1
            if len(sub_b) > min_len:
                for k in range(min_len, len(sub_b)):
                    diffs.append({"type": "added", "text": sub_b[k]})
                    count_added += 1

    has_changes = (count_added + count_deleted + count_modified) > 0
    warnings = [
        "Textual comparison performed via difflib on extracted text. This is an approximate textual diff, not a legal-grade semantic redline."
    ]

    is_markdown = str(format or "").lower().strip() in ("markdown", "md")
    result: dict[str, Any] = {
        "status": "ok",
        "path_a": path_a,
        "path_b": path_b,
        "format_a": Path(path_a).suffix.lower(),
        "format_b": Path(path_b).suffix.lower(),
        "has_changes": has_changes,
        "summary": {
            "added": count_added,
            "deleted": count_deleted,
            "modified": count_modified,
            "unchanged": count_unchanged,
        },
        "diffs": diffs,
        "warnings": warnings,
    }

    if is_markdown:
        lines = [
            f"# Document Diff: {Path(path_a).name} vs {Path(path_b).name}",
            "",
            f"**Summary:** {count_added} added, {count_deleted} deleted, {count_modified} modified, {count_unchanged} unchanged.",
            "",
            "### Changes:",
        ]
        if not diffs:
            lines.append("*(No changes detected)*")
        else:
            for d in diffs:
                dtype = d["type"]
                if dtype == "added":
                    lines.append(f"+ **[Added]** {d['text']}")
                elif dtype == "deleted":
                    lines.append(f"- **[Deleted]** {d['text']}")
                elif dtype == "modified":
                    lines.append(f"~ **[Modified]**\n  * **Old:** {d['old_text']}\n  * **New:** {d['new_text']}")
        lines.append("")
        lines.append(f"> **Warning:** {warnings[0]}")
        result["diff_markdown"] = "\n".join(lines).strip()

    return result
